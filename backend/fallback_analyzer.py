import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import sys
import json
import os
import re
import warnings

# Suppress non-critical matplotlib/seaborn warnings
warnings.filterwarnings('ignore')

# 1. Premium Dark Theme Styles matching Web App
plt.rcParams['figure.facecolor'] = '#12121e'
plt.rcParams['axes.facecolor'] = '#1a1a2e'
plt.rcParams['grid.color'] = '#2d2d44'
plt.rcParams['grid.linestyle'] = '--'
plt.rcParams['grid.linewidth'] = 0.5
plt.rcParams['text.color'] = '#e2e8f0'
plt.rcParams['axes.labelcolor'] = '#e2e8f0'
plt.rcParams['xtick.color'] = '#94a3b8'
plt.rcParams['ytick.color'] = '#94a3b8'
plt.rcParams['font.family'] = 'sans-serif'

COLORS = ['#8b5cf6', '#06b6d4', '#ec4899', '#10b981', '#f97316', '#a78bfa', '#22d3ee', '#f472b6', '#3b82f6', '#eab308']

def format_currency(val):
    try:
        val_float = float(val)
        return f"${val_float:,.2f}"
    except (ValueError, TypeError):
        return str(val)

def find_col(df, target_names):
    """
    Finds the first matching column in df (case-insensitive, substring/alias matching).
    """
    cols_lower = {str(c).lower().strip(): c for c in df.columns}
    for target in target_names:
        if target.lower() in cols_lower:
            return cols_lower[target.lower()]
    for target in target_names:
        for c_lower, orig in cols_lower.items():
            if target.lower() in c_lower:
                return orig
    return None

def load_any_file_as_df(file_path, sheet_name=None):
    """
    Universal multi-format parser for Excel (.xlsx, .xls), CSV (.csv, .tsv, .txt), 
    JSON (.json), Word Documents (.docx), and PDF files (.pdf).
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext in ['.xlsx', '.xls']:
        if sheet_name:
            return pd.read_excel(file_path, sheet_name=sheet_name)
        return pd.read_excel(file_path)
        
    elif ext in ['.csv', '.tsv', '.txt']:
        sep = '\t' if ext == '.tsv' else ','
        try:
            return pd.read_csv(file_path, sep=sep)
        except:
            return pd.read_csv(file_path, sep=r'\s+')
            
    elif ext == '.json':
        return pd.read_json(file_path)
        
    elif ext == '.docx':
        try:
            import docx
            doc = docx.Document(file_path)
            all_tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text.strip() for cell in row.cells])
                if len(table_data) > 1:
                    df_tbl = pd.DataFrame(table_data[1:], columns=table_data[0])
                    all_tables.append(df_tbl)
            if all_tables:
                return pd.concat(all_tables, ignore_index=True)
            
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            return pd.DataFrame({"Content": lines})
        except Exception as e:
            raise ValueError(f"Could not parse Word document: {str(e)}")
            
    elif ext == '.pdf':
        try:
            import pdfplumber
            tables = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_tables()
                    for tbl in extracted:
                        if len(tbl) > 1:
                            df_tbl = pd.DataFrame(tbl[1:], columns=tbl[0])
                            tables.append(df_tbl)
            if tables:
                return pd.concat(tables, ignore_index=True)
                
            lines = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    txt = page.extract_text()
                    if txt:
                        lines.extend([l.strip() for l in txt.split('\n') if l.strip()])
            return pd.DataFrame({"Content": lines})
        except Exception as e:
            raise ValueError(f"Could not parse PDF document: {str(e)}")
    else:
        return pd.read_csv(file_path)

def prepare_dataframe(df):
    df.columns = [str(c).strip() for c in df.columns]
    
    # 1. Customer Name column
    if 'Customer Name' not in df.columns:
        cust_col = find_col(df, ['customer name', 'customer', 'client', 'name', 'customer_name', 'client_name', 'user', 'email'])
        fn_col = find_col(df, ['first_name', 'firstname', 'given_name'])
        ln_col = find_col(df, ['last_name', 'lastname', 'surname', 'family_name'])
        
        if fn_col and ln_col:
            df['Customer Name'] = df[fn_col].astype(str) + ' ' + df[ln_col].astype(str)
        elif cust_col:
            df['Customer Name'] = df[cust_col].astype(str)
        elif fn_col:
            df['Customer Name'] = df[fn_col].astype(str)
        elif ln_col:
            df['Customer Name'] = df[ln_col].astype(str)
        else:
            df['Customer Name'] = 'Customer ' + (df.index + 1).astype(str)
            
    # 2. Invoice Amount column
    if 'Invoice Amount' not in df.columns:
        amt_col = find_col(df, ['invoice amount', 'invoice_amount', 'total_amount', 'total', 'amount', 'price', 'value', 'grand_total', 'cost', 'sales', 'revenue'])
        if amt_col:
            qty_col = find_col(df, ['qty', 'quantity', 'count'])
            df['Invoice Amount'] = pd.to_numeric(df[amt_col], errors='coerce').fillna(0)
            if qty_col and amt_col.lower() in ['price', 'unit_price', 'unit price']:
                df['Invoice Amount'] = df['Invoice Amount'] * pd.to_numeric(df[qty_col], errors='coerce').fillna(1)
        else:
            num_cols = df.select_dtypes(include=[np.number]).columns
            if len(num_cols) > 0:
                df['Invoice Amount'] = df[num_cols[0]].fillna(0)
            else:
                df['Invoice Amount'] = 0.0

    # 3. Status column
    if 'Status' not in df.columns:
        stat_col = find_col(df, ['status', 'payment_status', 'invoice_status', 'state'])
        if stat_col:
            df['Status'] = df[stat_col].astype(str)
        else:
            df['Status'] = 'Processed'

    # 4. Amount Paid & Outstanding Balance
    if 'Amount Paid' not in df.columns:
        paid_col = find_col(df, ['amount paid', 'amount_paid', 'paid_amount', 'paid'])
        if paid_col:
            df['Amount Paid'] = pd.to_numeric(df[paid_col], errors='coerce').fillna(0)
        else:
            df['Amount Paid'] = df.apply(
                lambda r: r['Invoice Amount'] if str(r['Status']).lower() in ['paid', 'completed', 'success'] else 0.0, 
                axis=1
            )
            
    if 'Outstanding Balance' not in df.columns:
        out_col = find_col(df, ['outstanding balance', 'outstanding_balance', 'balance', 'outstanding', 'unpaid', 'due_amount'])
        if out_col:
            df['Outstanding Balance'] = pd.to_numeric(df[out_col], errors='coerce').fillna(0)
        else:
            df['Outstanding Balance'] = (df['Invoice Amount'] - df['Amount Paid']).clip(lower=0)

    # 5. Overdue Days
    if 'Overdue Days' not in df.columns:
        od_col = find_col(df, ['overdue days', 'overdue_days', 'days_overdue', 'overdue', 'days_late'])
        if od_col:
            df['Overdue Days'] = pd.to_numeric(df[od_col], errors='coerce').fillna(0)
        else:
            date_col = find_col(df, ['invoice date', 'invoice_date', 'date', 'created_at', 'due date', 'due_date'])
            if date_col:
                try:
                    parsed_dates = pd.to_datetime(df[date_col], errors='coerce')
                    ref_date = parsed_dates.max()
                    if pd.notnull(ref_date):
                        days_diff = (ref_date - parsed_dates).dt.days
                        df['Overdue Days'] = days_diff.fillna(0).astype(int)
                    else:
                        df['Overdue Days'] = 0
                except:
                    df['Overdue Days'] = 0
            else:
                df['Overdue Days'] = 0

    # 6. Category column
    if 'Category' not in df.columns:
        cat_col = find_col(df, ['category', 'product_id', 'product', 'item', 'job', 'city', 'department', 'type'])
        if cat_col:
            df['Category'] = df[cat_col].astype(str)
        else:
            df['Category'] = 'General'
            
    # 7. Invoice Date
    if 'Invoice Date' not in df.columns:
        d_col = find_col(df, ['invoice date', 'invoice_date', 'date', 'created_at', 'timestamp'])
        if d_col:
            df['Invoice Date'] = df[d_col].astype(str)
        else:
            df['Invoice Date'] = '2025-01-01'

    return df

def run_analysis(file_path, sheet_name, query, chart_path):
    df = load_any_file_as_df(file_path, sheet_name)
    df = prepare_dataframe(df)
    
    q = query.lower()
    
    result = {
        "text_response": "",
        "table_data": None,
        "table_columns": None,
        "has_chart": False,
        "chart_type": None,
        "success": True,
        "error_message": None
    }
    
    def save_plot(chart_type):
        plt.tight_layout()
        plt.savefig(chart_path, dpi=150, facecolor='#12121e', edgecolor='none')
        plt.close('all')
        result["has_chart"] = True
        result["chart_type"] = chart_type

    try:
        city_col = find_col(df, ['city', 'location', 'town', 'region', 'address'])

        # CASE 1: Invoice Status Chart (Pie / Donut Chart)
        if ("status" in q or "paid" in q) and ("bar" in q or "chart" in q or "visual" in q or "distribution" in q or "breakdown" in q or "pie" in q):
            status_counts = df['Status'].value_counts()
            
            fig, ax = plt.subplots(figsize=(8, 8))
            pie_colors = (COLORS * ((len(status_counts) // len(COLORS)) + 1))[:len(status_counts)]
            ax.pie(status_counts.values, labels=status_counts.index, autopct='%1.1f%%', startangle=90, 
                   colors=pie_colors, textprops={'color': 'white', 'fontsize': 11},
                   wedgeprops=dict(width=0.4, edgecolor='#12121e'))
            ax.set_title("Invoice Count Distribution by Status", fontsize=14, color='white', weight='bold', pad=20)
                
            save_plot("pie")
            
            total = len(df)
            markdown = "### Invoice Status Summary & Distribution\nBreakdown of invoice statuses across the dataset:\n\n"
            tbl_data = []
            for status, count in status_counts.items():
                pct = (count / total) * 100 if total > 0 else 0
                amount = df[df['Status'] == status]['Invoice Amount'].sum()
                markdown += f"- **{status}**: {count} records ({pct:.1f}%) totaling **{format_currency(amount)}**\n"
                tbl_data.append({
                    "Status": str(status),
                    "Count": int(count),
                    "Percentage": f"{pct:.1f}%",
                    "Total Value": format_currency(amount)
                })
            
            result["text_response"] = markdown
            result["table_data"] = tbl_data
            result["table_columns"] = ["Status", "Count", "Percentage", "Total Value"]
            
        # CASE 2: Top Customers by Invoice Value
        elif "top" in q or "customer" in q or "client" in q or ("highest" in q and "value" in q):
            n_match = re.search(r'top\s+(\d+)', q)
            top_n = int(n_match.group(1)) if n_match else 10
            
            top_cust = df.groupby('Customer Name')['Invoice Amount'].sum().reset_index()
            top_cust = top_cust.sort_values(by='Invoice Amount', ascending=False).head(top_n)
            
            fig, ax = plt.subplots(figsize=(8, 8))
            pie_colors = (COLORS * ((len(top_cust) // len(COLORS)) + 1))[:len(top_cust)]
            ax.pie(top_cust['Invoice Amount'], labels=top_cust['Customer Name'], autopct='%1.1f%%', startangle=90,
                   colors=pie_colors, textprops={'color': 'white', 'fontsize': 10},
                   wedgeprops=dict(width=0.4, edgecolor='#12121e'))
            ax.set_title(f"Top {top_n} Customers Share by Invoice Value", fontsize=14, color='white', weight='bold', pad=20)
            
            save_plot("pie")
            
            markdown = f"### Top {top_n} Customers by Total Invoice Value\nTop customers ranked by cumulative invoice amounts:\n\n"
            tbl_data = []
            tbl_cols = ["Rank", "Customer Name", "Total Invoiced"]
            if city_col:
                tbl_cols.insert(2, "City")

            for rank, (idx, row) in enumerate(top_cust.iterrows(), 1):
                cust = row['Customer Name']
                amt = row['Invoice Amount']
                city_name = str(df[df['Customer Name'] == cust][city_col].iloc[0]) if city_col and len(df[df['Customer Name'] == cust]) > 0 else "-"
                
                if city_col and city_name != "-":
                    markdown += f"{rank}. **{cust}** ({city_name}): {format_currency(amt)}\n"
                else:
                    markdown += f"{rank}. **{cust}**: {format_currency(amt)}\n"

                row_dict = {
                    "Rank": rank,
                    "Customer Name": cust,
                    "Total Invoiced": format_currency(amt)
                }
                if city_col:
                    row_dict["City"] = city_name
                tbl_data.append(row_dict)
            
            result["text_response"] = markdown
            result["table_data"] = tbl_data
            result["table_columns"] = tbl_cols

        # CASE 3: Overdue Day Ranges Pie/Donut Chart
        elif "overdue" in q or "late" in q:
            def bucket_overdue(days):
                if days <= 0:
                    return "Current / Paid"
                elif days <= 30:
                    return "1-30 Days"
                elif days <= 60:
                    return "31-60 Days"
                elif days <= 90:
                    return "61-90 Days"
                else:
                    return "90+ Days Overdue"
            
            df['Overdue Range'] = df['Overdue Days'].apply(bucket_overdue)
            ranges = df['Overdue Range'].value_counts()
            
            fig, ax = plt.subplots(figsize=(8, 8))
            ax.pie(ranges.values, labels=ranges.index, autopct='%1.1f%%', startangle=90, 
                   colors=COLORS[:len(ranges)], textprops={'color': 'white', 'fontsize': 11},
                   wedgeprops=dict(width=0.4, edgecolor='#12121e'))
            ax.set_title("Overdue Distribution by Days Range", fontsize=14, color='white', weight='bold', pad=20)
            
            save_plot("pie")
            
            markdown = "### Overdue Days Breakdown\nCategorization of records by overdue day ranges:\n\n"
            tbl_data = []
            total_count = ranges.sum()
            for r_name, count in ranges.items():
                pct = (count / total_count) * 100 if total_count > 0 else 0
                markdown += f"- **{r_name}**: {count} records ({pct:.1f}%)\n"
                tbl_data.append({
                    "Overdue Range": str(r_name),
                    "Count": int(count),
                    "Percentage": f"{pct:.1f}%"
                })
                
            result["text_response"] = markdown
            result["table_data"] = tbl_data
            result["table_columns"] = ["Overdue Range", "Count", "Percentage"]

        # CASE 4: Calculate Totals (Invoiced, Paid, Outstanding)
        elif "total" in q or "calculate" in q or "sum" in q or "metrics" in q:
            total_invoiced = df['Invoice Amount'].sum()
            total_paid = df['Amount Paid'].sum()
            total_outstanding = df['Outstanding Balance'].sum()
            paid_pct = (total_paid / total_invoiced) * 100 if total_invoiced > 0 else 0
            outstanding_pct = (total_outstanding / total_invoiced) * 100 if total_invoiced > 0 else 0
            
            fig, ax = plt.subplots(figsize=(8, 8))
            vals = [max(total_paid, 0.01), max(total_outstanding, 0.01)]
            ax.pie(vals, labels=['Collected / Paid', 'Outstanding Balance'], 
                   autopct='%1.1f%%', startangle=90, colors=['#10b981', '#ef4444'],
                   textprops={'color': 'white', 'fontsize': 12}, wedgeprops=dict(width=0.4, edgecolor='#12121e'))
            ax.set_title("Total Collections vs. Outstanding Balance", fontsize=14, color='white', weight='bold', pad=20)
            save_plot("pie")
            
            markdown = f"### Financial Metrics Summary\n" \
                        f"Aggregated financial breakdown calculated from the dataset:\n\n" \
                        f"- **Total Invoiced Amount**: {format_currency(total_invoiced)}\n" \
                        f"- **Total Collected**: {format_currency(total_paid)} ({paid_pct:.1f}% collected)\n" \
                        f"- **Total Outstanding Balance**: {format_currency(total_outstanding)} ({outstanding_pct:.1f}% outstanding)\n"
            
            tbl_data = [
                {"Metric": "Total Invoiced", "Value": format_currency(total_invoiced), "Proportion": "100.0%"},
                {"Metric": "Collected / Paid", "Value": format_currency(total_paid), "Proportion": f"{paid_pct:.1f}%"},
                {"Metric": "Outstanding Balance", "Value": format_currency(total_outstanding), "Proportion": f"{outstanding_pct:.1f}%"}
            ]
            
            result["text_response"] = markdown
            result["table_data"] = tbl_data
            result["table_columns"] = ["Metric", "Value", "Proportion"]

        # CASE 5: Highest Outstanding Balances
        elif "outstanding" in q or "debt" in q or "unpaid" in q or "balance" in q:
            out_cust = df.groupby('Customer Name')['Outstanding Balance'].sum().reset_index()
            out_cust = out_cust.sort_values(by='Outstanding Balance', ascending=False).head(10)
            
            fig, ax = plt.subplots(figsize=(8, 8))
            pie_colors = (COLORS * ((len(out_cust) // len(COLORS)) + 1))[:len(out_cust)]
            ax.pie(out_cust['Outstanding Balance'], labels=out_cust['Customer Name'], autopct='%1.1f%%', startangle=90,
                   colors=pie_colors, textprops={'color': 'white', 'fontsize': 10},
                   wedgeprops=dict(width=0.4, edgecolor='#12121e'))
            ax.set_title("Top 10 Customers Share by Outstanding Debt", fontsize=14, color='white', weight='bold', pad=20)
            
            save_plot("pie")
            
            markdown = "### Customers with Highest Outstanding Balance\nTop customers holding unpaid balances:\n\n"
            tbl_data = []
            tbl_cols = ["Rank", "Customer Name", "Outstanding Debt"]
            if city_col:
                tbl_cols.insert(2, "City")

            for rank, (idx, row) in enumerate(out_cust.iterrows(), 1):
                cust = row['Customer Name']
                bal = row['Outstanding Balance']
                city_name = str(df[df['Customer Name'] == cust][city_col].iloc[0]) if city_col and len(df[df['Customer Name'] == cust]) > 0 else "-"
                
                if city_col and city_name != "-":
                    markdown += f"{rank}. **{cust}** ({city_name}): {format_currency(bal)}\n"
                else:
                    markdown += f"{rank}. **{cust}**: {format_currency(bal)}\n"

                row_dict = {
                    "Rank": rank,
                    "Customer Name": cust,
                    "Outstanding Debt": format_currency(bal)
                }
                if city_col:
                    row_dict["City"] = city_name
                tbl_data.append(row_dict)
            
            result["text_response"] = markdown
            result["table_data"] = tbl_data
            result["table_columns"] = tbl_cols

        # CASE 6: Category / Product / Job / City / Month Comparison (Pie Chart Default)
        else:
            cat_col_name = 'Category'
            if "city" in q and city_col:
                cat_col_name = city_col
            elif city_col and not ('Category' in df.columns and df['Category'].nunique() > 1):
                cat_col_name = city_col

            cat_data = df.groupby(cat_col_name)['Invoice Amount'].sum().reset_index()
            cat_data = cat_data.sort_values(by='Invoice Amount', ascending=False).head(12)
            
            fig, ax = plt.subplots(figsize=(8, 8))
            pie_colors = (COLORS * ((len(cat_data) // len(COLORS)) + 1))[:len(cat_data)]
            ax.pie(cat_data['Invoice Amount'], labels=cat_data[cat_col_name].astype(str), autopct='%1.1f%%',
                   startangle=90, colors=pie_colors, textprops={'color': 'white', 'fontsize': 10},
                   wedgeprops=dict(width=0.4, edgecolor='#12121e'))
            ax.set_title(f"Invoice Value Distribution across {cat_col_name}", fontsize=14, color='white', weight='bold', pad=20)
            
            save_plot("pie")
            
            markdown = f"### Invoice Value Breakdown by {cat_col_name}\nTotals grouped by {cat_col_name.lower()}:\n\n"
            tbl_data = []
            tbl_cols = ["Rank", cat_col_name, "Total Value"]
            if city_col and city_col != cat_col_name:
                tbl_cols.insert(2, "City")

            for rank, (idx, row) in enumerate(cat_data.iterrows(), 1):
                cat = row[cat_col_name]
                amt = row['Invoice Amount']
                markdown += f"{rank}. **{cat}**: {format_currency(amt)}\n"
                row_dict = {
                    "Rank": rank,
                    cat_col_name: str(cat),
                    "Total Value": format_currency(amt)
                }
                if city_col and city_col != cat_col_name:
                    city_name = str(df[df[cat_col_name] == cat][city_col].iloc[0]) if len(df[df[cat_col_name] == cat]) > 0 else "-"
                    row_dict["City"] = city_name
                tbl_data.append(row_dict)
            
            result["text_response"] = markdown
            result["table_data"] = tbl_data
            result["table_columns"] = tbl_cols

    except Exception as e:
        result["success"] = False
        result["error_message"] = str(e)
        result["text_response"] = f"An error occurred while running dataset analysis: {str(e)}"
        
    print(f"<RESULT_JSON>{json.dumps(result)}</RESULT_JSON>")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Missing config path"}))
        sys.exit(1)
        
    config_path = sys.argv[1]
    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)
        
    file_path = config["filePath"]
    chart_path = config["chartPath"]
    query = config["query"]
    sheet_name = config.get("sheetName", "Sheet1")
    
    run_analysis(file_path, sheet_name, query, chart_path)
