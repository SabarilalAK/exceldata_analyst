import pandas as pd
import json
import sys
import os

def load_any_file_as_df(file_path, sheet_name=None):
    """
    Universal multi-format parser for Excel (.xlsx, .xls), CSV (.csv, .tsv, .txt), 
    JSON (.json), Word Documents (.docx), and PDF files (.pdf).
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext in ['.xlsx', '.xls']:
        if sheet_name:
            return pd.read_excel(file_path, sheet_name=sheet_name)
        return pd.read_excel(file_path)
        
    elif ext in ['.csv', '.tsv', '.txt']:
        sep = '\t' if ext == '.tsv' else ','
        try:
            return pd.read_csv(file_path, sep=sep)
        except:
            return pd.read_csv(file_path, sep=r'\s+')
            
    elif ext == '.json':
        return pd.read_json(file_path)
        
    elif ext == '.docx':
        try:
            import docx
            doc = docx.Document(file_path)
            all_tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text.strip() for cell in row.cells])
                if len(table_data) > 1:
                    df_tbl = pd.DataFrame(table_data[1:], columns=table_data[0])
                    all_tables.append(df_tbl)
            if all_tables:
                return pd.concat(all_tables, ignore_index=True)
            
            # Fallback to paragraph text lines
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            return pd.DataFrame({"Content": lines})
        except Exception as e:
            raise ValueError(f"Could not parse Word document: {str(e)}")
            
    elif ext == '.pdf':
        try:
            import pdfplumber
            tables = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_tables()
                    for tbl in extracted:
                        if len(tbl) > 1:
                            df_tbl = pd.DataFrame(tbl[1:], columns=tbl[0])
                            tables.append(df_tbl)
            if tables:
                return pd.concat(tables, ignore_index=True)
                
            # Fallback to text lines extraction
            lines = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    txt = page.extract_text()
                    if txt:
                        lines.extend([l.strip() for l in txt.split('\n') if l.strip()])
            return pd.DataFrame({"Content": lines})
        except Exception as e:
            raise ValueError(f"Could not parse PDF document: {str(e)}")
    else:
        # Generic fallback
        return pd.read_csv(file_path)

def extract_metadata(file_path):
    try:
        if not os.path.exists(file_path):
            return {"error": "File does not exist"}
        
        ext = os.path.splitext(file_path)[1].lower()

        if ext in ['.xlsx', '.xls']:
            xl = pd.ExcelFile(file_path)
            sheets_data = {}
            for sheet in xl.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet)
                preview_df = df.head(10).copy()
                
                for col in preview_df.columns:
                    if pd.api.types.is_datetime64_any_dtype(preview_df[col]):
                        preview_df[col] = preview_df[col].dt.strftime('%Y-%m-%d')
                    elif preview_df[col].dtype == 'object':
                        try:
                            preview_df[col] = pd.to_datetime(preview_df[col]).dt.strftime('%Y-%m-%d')
                        except:
                            pass
                    preview_df[col] = preview_df[col].fillna('').replace([float('inf'), float('-inf')], '')
                
                dtypes = {col: str(df[col].dtype) for col in df.columns}
                
                sheets_data[sheet] = {
                    "columns": list(df.columns),
                    "dtypes": dtypes,
                    "rowCount": len(df),
                    "preview": preview_df.to_dict(orient="records")
                }
            
            return {
                "sheets": sheets_data,
                "defaultSheet": xl.sheet_names[0]
            }
        else:
            df = load_any_file_as_df(file_path)
            preview_df = df.head(10).copy()
            
            for col in preview_df.columns:
                if preview_df[col].dtype == 'object':
                    try:
                        preview_df[col] = pd.to_datetime(preview_df[col]).dt.strftime('%Y-%m-%d')
                    except:
                        pass
                elif pd.api.types.is_datetime64_any_dtype(preview_df[col]):
                    preview_df[col] = preview_df[col].dt.strftime('%Y-%m-%d')
                
                preview_df[col] = preview_df[col].fillna('').replace([float('inf'), float('-inf')], '')
            
            dtypes = {col: str(df[col].dtype) for col in df.columns}
            
            return {
                "sheets": {
                    "Sheet1": {
                        "columns": list(df.columns),
                        "dtypes": dtypes,
                        "rowCount": len(df),
                        "preview": preview_df.to_dict(orient="records")
                    }
                },
                "defaultSheet": "Sheet1"
            }
    except Exception as e:
        return {"error": str(e)}

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No file path provided"}))
        sys.exit(1)
    
    file_path = sys.argv[1]
    metadata = extract_metadata(file_path)
    print(json.dumps(metadata))
